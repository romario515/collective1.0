import os
import sys
import subprocess

from docx import Document
from PyQt5.QtWidgets import QApplication, QWidget, QPushButton,  QLabel, QCheckBox, QListWidget, QFileDialog
from pathlib import Path
import cv2
from ultralytics import YOLO
import torch
from transformers import BertTokenizer, BertForSequenceClassification
import json
from openpyxl import Workbook, load_workbook
from odf.opendocument import load, OpenDocumentSpreadsheet, OpenDocumentText
from odf.table import Table, TableRow, TableCell
from odf.text import P



thres_hat = 0
thres_vest = 0
hat_frame = ''
vest_frame = ''
founded = []
founded_vest = []
match_count = 2
previous_prediction = None
consecutive_count = 0
current_glob = ''
class win2(QWidget):

    def __init__(self):
        super().__init__()
        self.initUI()

    old_write = ["", 0]
    output_file1 = 'results/report.txt'

    def add_text_to_odt(self,file_path, text_to_add):
        # Попытка загрузить существующий ODT файл

        if os.path.exists(file_path):
            doc = load(file_path)
        else:
            # Если файл не существует или поврежден, создаем новый
            doc = OpenDocumentText()

        # Создаем новый параграф с текстом
        p = P(text=text_to_add)
        doc.text.addElement(p)

        # Сохраняем изменения в файле
        doc.save(file_path)

    def add_data_to_ods(self,file_path, data, sheet_name="Sheet1"):
        # Попытка загрузить существующий ODS файл
        print(file_path)
        if os.path.exists(file_path):
            print('hekllo')
            doc = load(file_path)
            # Поиск таблицы по имени
            print('hekllo')
            tables = doc.spreadsheet.getElementsByType(Table)
            table = None
            for tbl in tables:
                if tbl.getAttribute("name") == sheet_name:
                    table = tbl
                    break
            # Если таблица с заданным именем не найдена, создаем новую
            if not table:
                table = Table(name=sheet_name)
                doc.spreadsheet.addElement(table)
        else:
            # Если файл не существует или поврежден, создаем новый
            doc = OpenDocumentSpreadsheet()
            table = Table(name=sheet_name)
            doc.spreadsheet.addElement(table)

        # Добавление данных в таблицу
        for row_data in data:
            tr = TableRow()
            table.addElement(tr)
            for cell_data in row_data:
                tc = TableCell()
                tr.addElement(tc)
                text = P(text=str(cell_data))
                tc.addElement(text)

        # Сохраняем файл
        doc.save(file_path)

    def line_inserted(self):
        # Записываем последнюю полученную строку во все возможные файлы
        laststr=str(self.listWidget_f.item(self.listWidget_f.count()-1).text())

        with open(self.output_file1, 'a') as f:
            f.write(f"{laststr}\n")

        if os.path.exists('results/report.docx'):
            # Загрузка существующего документа
            doc = Document('results/report.docx')
        else:
            # Создание нового документа
            doc = Document()
            doc.save('results/report.docx')

        doc.add_paragraph(str(laststr))
        doc.save('results/report.docx')

        if os.path.exists('results/report.xlsx'):
            workbook = load_workbook(filename='results/report.xlsx')
            sheet = workbook['Sheet']
            if globals()['current_glob'] != laststr:
                sheet.append([globals()['current_glob'], str(laststr)])
            else:
                pass
            print('ENTERED')
            workbook.save(filename='results/report.xlsx')
            workbook.close()
        else:
            # Если файл не найден, создаем новый
            print('ENTERED_NEW')
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = 'Sheet'
            sheet.append(['Файл', 'Время'])
            workbook.save(filename='results/report.xlsx')
            workbook.close()
        print(laststr)
        self.add_data_to_ods('results/report.odf',[str(globals()['current_glob']), str(laststr)])
        self.add_text_to_odt('results/report.odt',str(laststr))

        print('check')
            # Добавление данных в новую строку



    def initUI(self):

        self.setGeometry(740, 100, 640, 480)
        self.setWindowTitle('Коллектив 1.0 - Найденное')
        self.listWidget_f = QListWidget(self)
        self.listWidget_f.resize(571, 400)
        self.listWidget_f.move(40, 35)
        self.listWidget_f.model().rowsInserted.connect(self.line_inserted)
        self.listWidget_f.itemSelectionChanged.connect(self.selectionChanged)


    def work(self):
        self.set_zero_glob()

        model1 = YOLO('results/best_peoples.pt')
        model2 = YOLO('results/best_train_f.pt')
        model3 = YOLO('results/best_rails_roma.pt')
        for x in range(ex.listWidget.count()):
            # Открытие видеофайла
            current = ex.listWidget.item(x).text()
            globals()['current_glob'] = current

            cap = cv2.VideoCapture(current)
            print(current)

            self.listWidget_f.addItem(str(current))
            print('after')
            # Получение ширины, высоты и FPS видео
            frame_width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            frame_height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            fps = cap.get(cv2.CAP_PROP_FPS)

            # Создание директории для сохранения результатов
            os.makedirs('dataset/results', exist_ok=True)
            # Загрузка токенизатора и модели для текстовой классификации
            tokenizer = BertTokenizer.from_pretrained(
                'results/')
            model = BertForSequenceClassification.from_pretrained(
                'results/')

            # Загрузка словаря меток
            with open('results/label_to_id.json',
                      'r') as f:
                label_to_id = json.load(f)

            # Обратный словарь для отображения меток
            id_to_label = {v: k for k, v in label_to_id.items()}

            # Переменная для настройки количества одинаковых предсказаний подряд
            consecutive_predictions_threshold = 4

            # Переменные для отслеживания предсказаний
            previous_prediction = None
            consecutive_count = 0

            def predict(text):
                inputs = tokenizer(text, return_tensors='pt', padding=True, truncation=True, max_length=512)
                outputs = model(**inputs)
                predictions = torch.nn.functional.softmax(outputs.logits, dim=-1)
                return predictions

            def process_frame(frame, models, model_names, timestamp, output_file):
                global previous_prediction, consecutive_count

            def save_results_to_file(model, results, model_name, output_file, timestamp):
                lines = []
                for result in results:
                    for bbox in result.boxes:
                        x1, y1, x2, y2 = bbox.xyxy[0]
                        width = x2 - x1
                        height = y2 - y1
                        confidence = bbox.conf[0] * 100  # Процент предсказания
                        class_id = bbox.cls[0]  # ID класса
                        class_name = model.names[int(class_id)]  # Название класса
                        lines.append(f"{model_name}|{class_name}|{confidence:.2f}|{x1}|{y1}|{width}|{height}")

                with open(output_file, 'a') as f:
                    f.write(f"{len(lines)}\n")  # Записываем количество строк
                    for line in lines:
                        f.write(f"{timestamp} {line}\n")  # Записываем каждую строку с временной меткой
                        print(f"{timestamp} {line}")

            def process_frame(frame, models, model_names, output_file, timestamp):
                all_results = []
                all_res_mas = []
                for model, model_name in zip(models, model_names):
                    results = model(frame)
                    for result in results:
                        for bbox in result.boxes:
                            x1, y1, x2, y2 = bbox.xyxy[0]
                            width = x2 - x1
                            height = y2 - y1
                            confidence = bbox.conf[0] * 100  # Процент предсказания
                            class_id = bbox.cls[0]  # ID класса
                            class_name = model.names[int(class_id)]  # Название класса
                            all_results.append(f"{model_name}|{class_name}|{confidence:.2f}|{x1}|{y1}|{width}|{height}")
                            all_res_mas.append(
                                [model_name, class_name, confidence, float(x1), float(y1), x2, y2, width, height,
                                 timestamp])
                            # Визуализация результатов
                            cv2.rectangle(frame, (int(x1), int(y1)), (int(x2), int(y2)), (255, 0, 0), 2)
                            cv2.putText(frame, f"{class_name} {confidence:.2f}%", (int(x1), int(y1) - 10),
                                        cv2.FONT_HERSHEY_SIMPLEX,
                                        0.9, (255, 0, 0), 2)

                            # Подготовка текста для классификации
                input_text = f"{len(all_results)}\n" + "\n".join(all_results)

                # Предсказание на основе результатов детекции
                predictions = predict(input_text)
                predicted_label_id = torch.argmax(predictions, dim=1).item()
                predicted_label = id_to_label[predicted_label_id]

                if predicted_label != "not":
                    if predicted_label != "blizko1":
                        if predicted_label != "Kaska2":
                            print("До", globals()['previous_prediction'], "После", predicted_label)
                            if predicted_label == globals()['previous_prediction']:
                                globals()['consecutive_count'] += 1
                            else:
                                globals()['previous_prediction'] = predicted_label
                                globals()['consecutive_count'] = 1

                            if globals()['consecutive_count'] >= consecutive_predictions_threshold:
                                out = 'не понятно'
                                if not (self.old_write[0] == predicted_label and timestamp - self.old_write[1] < 10):
                                    if (predicted_label == "zalez1"): out = ' Работа в подвагонном пространстве'
                                    if (predicted_label == "slez1"): out = ' Сход с подвижного состава'
                                    with open(output_file, 'a') as f:
                                        f.write(f"{timestamp:.4f} {out}\n")
                                        self.listWidget_f.addItem(str(timestamp)+ out)
                                self.old_write[1] = float(timestamp)
                                self.old_write[0] = predicted_label
                                print(self.old_write[0], self.old_write[1], "==========")
                                globals()['consecutive_count'] = 0
                else:
                    globals()['previous_prediction'] = predicted_label
                    globals()['consecutive_count'] = 1
                # Вывод предсказаний
                print(f"Predicted label: {predicted_label}")

                # Вывод вероятностей для всех классов
                for i, prob in enumerate(predictions[0]):
                    label = id_to_label[i]
                    print(f"{label}: {prob.item() * 100:.2f}%")
                # Сохранение результатов в текстовый файл
                with open(output_file, 'a') as f:
                    f.write(f"{len(all_results)}\n")  # Записываем количество строк
                    for line in all_results:
                        f.write(f"{line}\n")  # Записываем каждую строку с временной меткой
                        print(f"{timestamp} {line}")
                person_list = [i for i, j in enumerate(all_res_mas) if j[1] == 'Person']
                nohat_list = [i for i, j in enumerate(all_res_mas) if j[1] == 'NO-Hardhat']
                novest_list = [i for i, j in enumerate(all_res_mas) if j[1] == 'NO-Safety Vest']

                if len(person_list) > 0 and len(nohat_list) > 0:
                    for x in person_list:
                        for j in nohat_list:
                            if all_res_mas[x][3] <= all_res_mas[j][3] <= all_res_mas[x][5]:
                                if all_res_mas[x][4] <= all_res_mas[j][6] <= all_res_mas[x][6]:
                                    if globals()['thres_hat'] < 1:
                                        globals()['hat_frame'] = round(all_res_mas[x][9], 1)
                                    if 0 <= globals()['thres_hat'] <= globals()['match_count']: globals()[
                                        'thres_hat'] += 1
                                    print('Match Hat!')
                            else:
                                if 0 <= globals()['thres_hat'] <= globals()['match_count']: globals()['thres_hat'] -= 1
                if len(person_list) > 0 and len(novest_list) > 0:
                    for x in person_list:
                        for j in novest_list:
                            if all_res_mas[x][3] <= all_res_mas[j][3] <= all_res_mas[x][5]:
                                if all_res_mas[x][4] <= all_res_mas[j][6] <= all_res_mas[x][6]:
                                    if globals()['thres_vest'] < 1:
                                        globals()['vest_frame'] = round(all_res_mas[x][9], 1)
                                    if 0 <= globals()['thres_vest'] <= globals()['match_count']: globals()[
                                        'thres_vest'] += 1
                                    print('Match Vest!')
                            else:
                                if 0 <= globals()['thres_vest'] <= globals()['match_count']: globals()[
                                    'thres_vest'] -= 1

                print(person_list)
                print(nohat_list)
                return frame

            # Основной цикл обработки
            frame_count = 0  # Счетчик кадров
            paused = False  # Флаг паузы
            desired_fps = 4  # Укажите здесь желаемую частоту кадров
            frame_interval = int(fps / desired_fps)  # Интервал для чтения кадров

            while cap.isOpened():
                if not paused:
                    ret, frame = cap.read()
                    if not ret:
                        break
                    ts1 = ''
                    # Обработка только каждого n-го кадра для регулирования частоты кадров
                    if frame_count % frame_interval == 0:
                        timestamp = frame_count / fps  # Временная метка в секундах
                        ts1=timestamp
                        frame = process_frame(frame, [model1, model2, model3], ['Model 1', 'Model 2', 'Model 3'],
                                              'results/predictions.txt', timestamp)
                    if thres_hat > match_count:
                        if hat_frame not in founded:
                            if len(founded) < 1:
                                founded.append(hat_frame)
                                print(hat_frame)
                                self.listWidget_f.addItem(str(ts1) + ' - Отсутствует каска')
                            else:
                                if round(float(hat_frame) - founded[-1]) > 10:
                                    self.listWidget_f.addItem(str(ts1) + ' - Отсутствует каска')
                                    print(hat_frame)
                        # тут вывод на окно текущего видео, и в файлы
                    if thres_vest > match_count:
                        if vest_frame not in founded_vest:
                            if len(founded_vest) < 1:
                                founded_vest.append(vest_frame)
                                self.listWidget_f.addItem(str(ts1)+' - Отсутствует жилет')
                            else:
                                if round(float(vest_frame) - founded_vest[-1]) > 10:
                                    founded_vest.append(vest_frame)
                                    self.listWidget_f.addItem(str(ts1)+' - Отсутствует жилет')
                                    #print('Next in')
                    # Отображение кадра
                    cv2.imshow('Frame', frame)

                    frame_count += 1

                # Обработка нажатия клавиш
                key = cv2.waitKey(1) & 0xFF
                if key == ord('q'):
                    break
                elif key == ord('p'):
                    paused = not paused

            # Освобождение ресурсов
            cap.release()
            cv2.destroyAllWindows()


    def set_zero_glob(self):
        globals()['thres_hat']=0
        globals()['thres_vest'] = 0
        globals()['hat_frame'] = ''
        globals()['vest_frame'] = ''
        globals()['founded'] = []
        globals()['founded_vest'] = []

    def selectionChanged(self):
        print('try')

        list=self.listWidget_f.currentItem().text().split()
        print(list[0])
        print(r'start wmplayer ' + str(current_glob) + ' /play /start ' + str(round(float(list[0])*1000)))
        #os.system(f'start wmplayer ' + str(current_glob) + ' /play /start ' + str(round(float(list[0])*1000)))
        subprocess.run(r'start wmplayer "' + str(current_glob) + '" /play /start' + str(round(float(list[0])*1000)), shell=True)
        print(str)




class win(QWidget):
    def __init__(self):
        super().__init__()
        self.initUI()
        self.items =[]

    def initUI(self):
        self.mode = {'w': 0, 'e': 0, 'i': 0}
        self.setGeometry(100, 100, 640, 480)
        self.setWindowTitle('Коллектив 1.0')

        # создаем кнопку
        self.btn = QPushButton('Загрузить видео', self)
        self.btn_start = QPushButton('Найти нарушения', self)
        self.btn_clear = QPushButton('Очистить список', self)
        # изменяем рамер кнопки
        self.btn.resize(150, 50)
        self.btn_start.resize(150, 50)
        self.btn_clear.resize(150, 50)
        # размещаем кпопку на родительском виджете
        self.btn.move(40, 370)
        self.btn_start.move(460, 370)
        self.btn_clear.move(240, 370)
        # определяем обработчик нажатия кнопик
        self.btn.clicked.connect(self.get_files)
        self.btn_clear.clicked.connect(self.clear_list)
        self.listWidget = QListWidget(self)
        self.listWidget.resize(571, 291)
        self.listWidget.move(40, 35)
        self.btn_start.clicked.connect(self.show_window2)
        self.btn.setStyleSheet("background-color: #4CAF50; /* Green */\n"
                                        "        border: none;\n"
                                        "        color: white;\n"
                                        #"        padding: 15px 32px;\n"
                                        "        text-align: center;\n"
                                        "        text-decoration: none;\n"
                                        #"        display: inline-block;\n"
                                        "        font-size: 16px;")
        self.btn_start.setStyleSheet("background-color: #e21a1a; /* Green */\n"
                                        "        border: none;\n"
                                        "        color: white;\n"
                                       # "        padding: 15px 32px;\n"
                                        "        text-align: center;\n"
                                        "        text-decoration: none;\n"
                                        #"        display: inline-block;\n"
                                        "        font-size: 16px;")
        self.label = QLabel(self)
        #self.label.setFrameStyle(QFrame.Panel | QFrame.Sunken)
        self.label.setText("Формат вывода:")
        self.label.move(350, 340)

        self.check1 = QCheckBox("Word", self)
        self.check1.stateChanged.connect(self.check_1)
        self.check1.move(460, 340)

        self.check2 = QCheckBox("Exel",self)
        self.check2.stateChanged.connect(self.check_2)
        self.check2.move(515, 340)

        self.check3 = QCheckBox("Вывод данных в ods, odf", self)
        self.check3.stateChanged.connect(self.check_3)
        self.check3.move(100, 340)


        self.label_vids = QLabel(self)
        # self.label_vids.setFrameStyle(QFrame.Panel | QFrame.Sunken)
        self.label_vids.setText("Выбранные файлы:")
        self.label_vids.move(40, 10)
        # listWidgetItem = QListWidgetItem("GeeksForGeeks")
        # self.listWidget.addItem(listWidgetItem)

    def clos(self): # метод закрывает окно
        self.close()

    def clear_list(self):
        self.listWidget.clear()
        self.listWidget.show()

    def get_files(self):
        dialog = QFileDialog()
        dialog.setDirectory(r'C:\Users')
        dialog.setFileMode(QFileDialog.FileMode.ExistingFiles)
        # dialog.setNameFilter("Images (*.png *.jpg)")
        dialog.setViewMode(QFileDialog.ViewMode.List)
        if dialog.exec():
            filenames = dialog.selectedFiles()
            if filenames:
                for filename in filenames:
                    if self.items.count(str(Path(filename))) > 0:
                        print(self.items.count(str(filename)))
                    else:
                        self.items.append(str(Path(filename)))
                        self.listWidget.addItem(str(Path(filename)))

    def check_1(self, checked):
        if checked:
            self.mode['w'] = 1
        else:
            self.mode['w'] = 0
            self.show()

    def check_2(self, checked):
        if checked:
            self.mode['e'] = 1
        else:
            self.mode['e'] = 0
            self.show()

    def check_3(self, checked):
        if checked:
            self.mode['i'] = 1
        else:
            self.mode['i'] = 0
            self.show()

    def show_window2(self):
        self.w2 = win2()
        self.w2.show()
        if os.path.exists('results/report.docx'):
            # Загрузка существующего документа
            doc = Document('results/report.docx')
            for element in doc.element.body:
                doc.element.body.remove(element)
            doc.save('results/report.docx')
        if os.path.exists('results/report.xlsx'):
            workbook = load_workbook(filename='results/report.xlsx')
            for sheet in workbook.sheetnames:
                del workbook[sheet]
            sheet = workbook.create_sheet('Sheet')
            sheet.append(['Файл', 'Время'])
            workbook.save(filename='results/report.xlsx')
            workbook.close()
        try:
            doc_o = OpenDocumentSpreadsheet(filename='results/report.ods')
            # Удаляем все существующие таблицы
            for table in doc_o.spreadsheet.getElementsByType(Table):
                doc_o.spreadsheet.removeChild(table)
        except Exception:
            # Если файл не существует или поврежден, создаем новый
            doc_o = OpenDocumentSpreadsheet()
        table = Table(name="Sheet1")
        doc_o.spreadsheet.addElement(table)

        self.w2.work()



    def closeEvent(self, event):
        ex.closeAllWindows()





app = QApplication(sys.argv)
ex = win()
ex.show()
sys.exit(app.exec())