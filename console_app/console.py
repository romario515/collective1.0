import os
import cv2
import torch
import json
from ultralytics import YOLO
from transformers import BertTokenizer, BertForSequenceClassification
import pandas as pd
from pathlib import Path
import sys
import subprocess
from docx import Document
from openpyxl import Workbook, load_workbook
from odf.opendocument import load, OpenDocumentSpreadsheet, OpenDocumentText
from odf.table import Table, TableRow, TableCell
from odf.text import P
old_write=["",0]
# Получение пути к текущему скрипту
script_path = os.path.abspath(__file__)

# Получение пути к директории
script_dir = os.path.dirname(script_path)
print(script_dir)

# Настройки
videos_folder = os.path.join(script_dir, 'videos_folder')  # Укажите путь к папке с видео
output_file1 = os.path.join(script_dir, 'report.txt')
output_doc = os.path.join(script_dir, 'results', 'report.docx')
output_xlsx = os.path.join(script_dir, 'results', 'report.xlsx')
consecutive_predictions_threshold = 3
match_count = 2

# Глобальные переменные
thres_hat = 0
thres_vest = 0
hat_frame = ''
vest_frame = ''
founded = []
founded_vest = []
previous_prediction = None
consecutive_count = 0
current_glob = ''

# Загрузка моделей YOLO
model1 = YOLO(os.path.join(script_dir, 'best_peoples.pt'))
model2 = YOLO(os.path.join(script_dir, 'best_train_f.pt'))
model3 = YOLO(os.path.join(script_dir, 'best_rails_roma.pt'))

# Загрузка токенизатора и модели для текстовой классификации
tokenizer = BertTokenizer.from_pretrained(os.path.join(script_dir, 'results'))
model = BertForSequenceClassification.from_pretrained(os.path.join(script_dir, 'results'))

# Загрузка словаря меток
with open(os.path.join(script_dir, 'results', 'label_to_id.json'), 'r') as f:
    label_to_id = json.load(f)

# Обратный словарь для отображения меток
id_to_label = {v: k for k, v in label_to_id.items()}

# Инициализация документа Word и Excel
# mydoc = docx.Document()

# def line_inserted(laststr, current_video):
#    with open(output_file1, 'a') as f:
#        f.write(f"{laststr}\n")  # Записываем последнюю полученную строку
#    mydoc.add_paragraph(laststr)
#    mydoc.save(output_doc)
#    if laststr != globals()['current_glob']:
#       df = pd.DataFrame({'Файл': [globals()['current_glob']], 'Время нарушения': [laststr]})
#       df.to_excel(output_xlsx)

def predict(text):
    inputs = tokenizer(text, return_tensors='pt', padding=True, truncation=True, max_length=512)
    outputs = model(**inputs)
    predictions = torch.nn.functional.softmax(outputs.logits, dim=-1)
    return predictions

def add_text_to_odt(file_path, text_to_add):
        # Попытка загрузить существующий ODT файл

        if os.path.exists(file_path):
            doc = load(file_path)
        else:
            # Если файл не существует или поврежден, создаем новый
            doc = OpenDocumentText()

        # Создаем новый параграф с текстом
        p = P(text=text_to_add)
        doc.text.addElement(p)

        # Сохраняем изменения в файле
        doc.save(file_path)

def add_data_to_ods(file_path, data):
    # Попытка загрузить существующий ODS файл
    sheet_name="Sheet1"
    print(file_path)
    if os.path.exists(file_path):
        print('hekllo')
        doc = load(file_path)
        # Поиск таблицы по имени
        print('hekllo')
        tables = doc.spreadsheet.getElementsByType(Table)
        table = None
        for tbl in tables:
            if tbl.getAttribute("name") == sheet_name:
                table = tbl
                break
        # Если таблица с заданным именем не найдена, создаем новую
        if not table:
            table = Table(name=sheet_name)
            doc.spreadsheet.addElement(table)
    else:
        # Если файл не существует или поврежден, создаем новый
        doc = OpenDocumentSpreadsheet()
        table = Table(name=sheet_name)
        doc.spreadsheet.addElement(table)

    # Добавление данных в таблицу
    print('before_data')
    for row_data in data:
        tr = TableRow()
        table.addElement(tr)
        for cell_data in row_data:
            tc = TableCell()
            tr.addElement(tc)
            text = P(text=str(cell_data))
            tc.addElement(text)
    print('after_data')
    # Сохраняем файл
    doc.save(file_path)

def line_inserted_doc(file_path,string_data):
    # Записываем последнюю полученную строку во все возможные файлы
    laststr=string_data
    if os.path.exists(file_path):
        # Загрузка существующего документа
        doc = Document(file_path)
    else:
        # Создание нового документа
        doc = Document()
        doc.save(file_path)

    doc.add_paragraph(laststr)
    doc.save(file_path)

def line_inserted_xls(file_path, data):
    if os.path.exists(file_path):
        workbook = load_workbook(filename=file_path)
        sheet = workbook['Sheet']
        sheet.append([str(globals()['current_glob']), str(data)])
        workbook.save(filename=file_path)
        workbook.close()
    else:
        # Если файл не найден, создаем новый
        print('ENTERED_NEW')
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = 'Sheet'
        sheet.append(['Файл', 'Время'])
        workbook.save(filename=file_path)
        workbook.close()

def process_frame(frame, models, model_names, timestamp, output_file):
    global previous_prediction, consecutive_count
    all_results = []
    all_res_mas = []
    for model, model_name in zip(models, model_names):
        results = model(frame)
        for result in results:
            for bbox in result.boxes:
                x1, y1, x2, y2 = bbox.xyxy[0]
                width = x2 - x1
                height = y2 - y1
                confidence = bbox.conf[0] * 100  # Процент предсказания
                class_id = bbox.cls[0]  # ID класса
                class_name = model.names[int(class_id)]  # Название класса
                all_results.append(f"{model_name}|{class_name}|{confidence:.2f}|{x1}|{y1}|{width}|{height}")
                print(timestamp,'---------------------------')
                all_res_mas.append([model_name, class_name, confidence, float(x1), float(y1), x2, y2, width, height, timestamp])
                # Визуализация результатов
                #cv2.rectangle(frame, (int(x1), int(y1)), (int(x2), int(y2)), (255, 0, 0), 2)
                #cv2.putText(frame, f"{class_name} {confidence:.2f}%", (int(x1), int(y1) - 10), cv2.FONT_HERSHEY_SIMPLEX, 0.9, (255, 0, 0), 2)

    # Подготовка текста для классификации
    input_text = f"{len(all_results)}\n" + "\n".join(all_results)

    # Предсказание на основе результатов детекции
    predictions = predict(input_text)
    predicted_label_id = torch.argmax(predictions, dim=1).item()
    predicted_label = id_to_label[predicted_label_id]

    if predicted_label != "not":
        if predicted_label != "blizko1":
            print("До",previous_prediction,"После",predicted_label )
            if predicted_label == previous_prediction:
                consecutive_count += 1
            else:
                previous_prediction = predicted_label
                consecutive_count = 1
            global old_write
            if consecutive_count >= consecutive_predictions_threshold:
                if not (old_write[0]==predicted_label and timestamp-old_write[1]<10):
                    with open(output_file, 'a') as f:
                        out='не понятно'
                        if(predicted_label == "zalez1"): out='Работа в подвагонном пространстве'
                        if(predicted_label == "slez1"): out='Сход с подвижного состава'
                        
                        f.write(f"{timestamp:.4f} {out}\n")
                        #add_data_to_ods('results/report.odf',[str(globals()['current_glob']), str(timestamp)])
                        #add_text_to_odt('results/report.odt',timestamp)
                        line_inserted_doc(output_doc, str(timestamp))
                        line_inserted_xls(output_xlsx,[globals()['current_glob'], timestamp])

                old_write[1]=float(timestamp)
                old_write[0] = predicted_label
                print(old_write[0],old_write[1],"==========")
                consecutive_count = 0
    else:
        previous_prediction = predicted_label
        consecutive_count = 1
    # Вывод предсказаний
    print(f"Predicted label: {predicted_label}")

    # Вывод вероятностей для всех классов
    for i, prob in enumerate(predictions[0]):
        label = id_to_label[i]
        #print(f"{label}: {prob.item() * 100:.2f}%")

    # Сохранение результатов в текстовый файл
    # with open(output_file, 'a') as f:
    #print(f"{len(all_results)}\n")  # Записываем количество строк
    #for line in all_results:
        #print(f"{line}\n")  # Записываем каждую строку с временной меткой

    person_list = [i for i, j in enumerate(all_res_mas) if j[1] == 'Person']
    nohat_list = [i for i, j in enumerate(all_res_mas) if j[1] == 'NO-Hardhat']
    novest_list = [i for i, j in enumerate(all_res_mas) if j[1] == 'NO-Safety Vest']

    if len(person_list) > 0 and len(nohat_list) > 0:
        for x in person_list:
            for j in nohat_list:
                if all_res_mas[x][3] <= all_res_mas[j][3] <= all_res_mas[x][5]:
                    if all_res_mas[x][4] <= all_res_mas[j][6] <= all_res_mas[x][6]:
                        if globals()['thres_hat'] < 1:
                            globals()['hat_frame'] = round(all_res_mas[x][9]*10, 1)
                            print(hat_frame,'!!!!!!!!!!!!!')
                        if 0 <= globals()['thres_hat'] <= globals()['match_count']:
                            globals()['thres_hat'] += 1
                        print('Match Hat!')
                else:
                    if 0 <= globals()['thres_hat'] <= globals()['match_count']:
                        globals()['thres_hat'] -= 1

    if len(person_list) > 0 and len(novest_list) > 0:
        for x in person_list:
            for j in novest_list:
                if all_res_mas[x][3] <= all_res_mas[j][3] <= all_res_mas[x][5]:
                    if all_res_mas[x][4] <= all_res_mas[j][6] <= all_res_mas[x][6]:
                        if globals()['thres_vest'] < 1:
                            globals()['vest_frame'] = round(all_res_mas[x][9], 1)
                        if 0 <= globals()['thres_vest'] <= globals()['match_count']:
                            globals()['thres_vest'] += 1
                        print('Match Vest!')
                else:
                    if 0 <= globals()['thres_vest'] <= globals()['match_count']:
                        globals()['thres_vest'] -= 1

    return frame

def set_zero_glob():
    globals()['thres_hat'] = 0
    globals()['thres_vest'] = 0
    globals()['hat_frame'] = ''
    globals()['vest_frame'] = ''
    globals()['founded'] = []
    globals()['founded_vest'] = []

def process_videos_in_folder(videos_folder):
    set_zero_glob()

    video_files = [str(p) for p in Path(videos_folder).glob("*.mp4")]

    for current in video_files:
        globals()['current_glob'] = current
        cap = cv2.VideoCapture(current)
        #print(f"Processing video: {current}")

        # Получение ширины, высоты и FPS видео
        frame_width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
        frame_height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
        fps = cap.get(cv2.CAP_PROP_FPS)

        frame_count = 0  # Счетчик кадров
        desired_fps = 4  # Укажите здесь желаемую частоту кадров
        frame_interval = int(fps / desired_fps)  # Интервал для чтения кадров
    
        while cap.isOpened():
            ret, frame = cap.read()
            if not ret:
                break
            st1=''
            # Обработка только каждого n-го кадра для регулирования частоты кадров
            if frame_count % frame_interval == 0:
                timestamp = frame_count / fps  # Временная метка в секундах
                st1=timestamp
                frame = process_frame(frame, [model1, model2, model3], ['Model 1', 'Model 2', 'Model 3'], timestamp, output_file1)
            
            if thres_hat > match_count:
                if hat_frame not in founded:
                    if len(founded) < 1:
                        founded.append(hat_frame)
                        #print(hat_frame)
                        print(hat_frame,' - Отсутствует каска')
                        with open(output_file1, 'a') as f:
                            f.write(str(st1) + ' - Отсутствует каска'+"\n")
                            add_data_to_ods('results/report.odf',[str(globals()['current_glob']), str(st1)])
                            add_text_to_odt('results/report.odt',str(st1))
                            line_inserted_doc(output_doc,  str(st1))
                            line_inserted_xls(output_xlsx,[str(globals()['current_glob']), str(st1)])
                    else:
                        if round(float(hat_frame) - founded[-1]) > 10:
                            print(str(st1) + ' - Отсутствует каска')
                            with open(output_file1, 'a') as f:
                                f.write(str(st1) + ' - Отсутствует каска'+"\n")
                            #print(hat_frame)
                # тут вывод на окно текущего видео, и в файлы
            if thres_vest > match_count:
                if vest_frame not in founded_vest:
                    if len(founded_vest) < 1:
                        founded_vest.append(vest_frame)
                        print(str(vest_frame)+' - Отсутствует жилет')
                        with open(output_file1, 'a') as f:
                            f.write(str(st1)+' - Отсутствует жилет'+"\n")
                    else:
                        if round(float(vest_frame) - founded_vest[-1]) > 10:
                            founded_vest.append(vest_frame)
                            print(str(vest_frame)+' - Отсутствует жилет')
                            with open(output_file1, 'a') as f:
                                f.write(str(st1)+' - Отсутствует жилет'+"\n")
                            #print('Next in')
            
            frame_count += 1

        cap.release()

if __name__ == "__main__":
    process_videos_in_folder(videos_folder)
